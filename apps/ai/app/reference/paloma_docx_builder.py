"""Constructor de documentos Word (.docx) con identidad visual de marca.

Todo el "componente estético" vive aquí (no en el LLM): el modelo solo aporta
contenido estructurado y este módulo lo convierte en un .docx elegante aplicando
una paleta, tipografías, portada, encabezado/pie y estilos consistentes.

Dos entradas públicas:
  - build_docx(spec, agente, ...) -> bytes      # universal, cualquier agente
  - build_tutela(payload, agente, ...) -> bytes # especializado, Decreto 2591/1991

Sin dependencias nuevas: usa python-docx (ya en requirements.txt).
"""
from __future__ import annotations

import io
import logging
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor, Inches

logger = logging.getLogger("paloma.documents")

THEME_VERSION = "2026.06.1"

# ── Paleta de campaña (oficial / presidencial) ─────────────────────────────
PRIMARY = "1A2A4F"   # azul noche — títulos y encabezados
ACCENT = "B08D43"    # dorado — reglas de acento y detalles
LIGHT = "EEF1F6"     # tinte claro — bandas de tabla
MUTED = "6B7280"     # gris — metadatos y pie
TEXT = "1A1A1A"      # negro suave — cuerpo

HEADING_FONT = "Georgia"   # serif elegante para títulos
BODY_FONT = "Calibri"      # sans legible para cuerpo

_MONTHS_ES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _fecha_larga(dt: datetime) -> str:
    return f"{dt.day} de {_MONTHS_ES[dt.month]} de {dt.year}"


# ── Helpers XML de bajo nivel ──────────────────────────────────────────────

def _shade(element, hex_color: str) -> None:
    """Aplica relleno de fondo (w:shd) a un párrafo (pPr) o celda (tcPr)."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pr.append(shd)


def _cell_shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bottom_border(paragraph, hex_color: str, size: int = 12) -> None:
    """Regla de acento bajo un párrafo (borde inferior)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _page_number_field(paragraph) -> None:
    """Inserta el campo dinámico PAGE en un párrafo."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_watermark(section, text: str) -> None:
    """Marca de agua diagonal en el encabezado (defensa adicional; el pie es la
    garantía dura). Best-effort: si el render VML falla, no rompe el documento."""
    try:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        xml = (
            '<w:r %s>'
            '<w:rPr><w:noProof/></w:rPr>'
            '<w:pict>'
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'id="DocWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;width:520pt;'
            'height:110pt;rotation:-45;z-index:-251653120;mso-position-horizontal:center;'
            'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
            'mso-position-vertical-relative:margin" fillcolor="#D9D9D9" stroked="f">'
            '<v:textpath style="font-family:&quot;Arial&quot;;font-size:1pt" '
            'string="%s"/>'
            '</v:shape>'
            '</w:pict>'
            '</w:r>' % (nsdecls("w"), text)
        )
        run._r.addnext(parse_xml(xml))
    except Exception as e:  # pragma: no cover - defensa
        logger.warning(f"Watermark no aplicada: {e}")


# ── Tema (estilos de marca) ────────────────────────────────────────────────

def _apply_theme(doc: Document) -> None:
    # Cuerpo
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Títulos
    for name, size, color in (
        ("Title", 30, PRIMARY),
        ("Heading 1", 17, PRIMARY),
        ("Heading 2", 14, PRIMARY),
        ("Heading 3", 12, ACCENT),
    ):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = HEADING_FONT
        st.font.size = Pt(size)
        st.font.bold = name != "Title"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(14 if name != "Title" else 0)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _build_cover(doc: Document, titulo: str, subtitulo: str, agente: str,
                 tipo: str, fecha: str) -> None:
    # Banda de acento superior
    band = doc.add_paragraph()
    band.paragraph_format.space_after = Pt(2)
    _shade(band._p, ACCENT)
    band.add_run(" ").font.size = Pt(4)

    doc.add_paragraph().add_run().add_break()  # aire

    t = doc.add_paragraph(style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.add_run(titulo)
    _bottom_border(t, ACCENT, size=18)

    if subtitulo:
        s = doc.add_paragraph()
        sr = s.add_run(subtitulo)
        sr.font.name = HEADING_FONT
        sr.font.size = Pt(14)
        sr.font.italic = True
        sr.font.color.rgb = RGBColor.from_string(MUTED)

    for _ in range(3):
        doc.add_paragraph()

    # Bloque de metadatos
    for label, value in (
        ("Preparado por", agente),
        ("Tipo de documento", tipo.capitalize()),
        ("Fecha", fecha),
    ):
        if not value:
            continue
        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        lr = meta.add_run(f"{label}:  ")
        lr.font.size = Pt(10)
        lr.font.bold = True
        lr.font.color.rgb = RGBColor.from_string(PRIMARY)
        vr = meta.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = RGBColor.from_string(TEXT)

    # Salto de página tras la portada
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _build_footer(doc: Document, agente: str, fecha: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    # Tab derecho al ancho útil de la página
    usable = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run(f"Generado por {agente} · PipeOs · {fecha}")
    left.font.size = Pt(8)
    left.font.color.rgb = RGBColor.from_string(MUTED)
    p.add_run("\t")
    pg = p.add_run("Página ")
    pg.font.size = Pt(8)
    pg.font.color.rgb = RGBColor.from_string(MUTED)
    _page_number_field(p)


def _build_header(doc: Document, wordmark: str = "PipeOs") -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(wordmark)
    r.font.name = HEADING_FONT
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(PRIMARY)


# ── Render de secciones (camino universal) ─────────────────────────────────

def _normalize_filas(filas) -> list:
    """Coacciona variantes de forma del modelo a lista-de-listas-de-strings.

    El modelo a veces emite filas como dicts (valores), strings sueltas, o None.
    Sin esto, esas formas provocaban error o se perdían filas."""
    out: list = []
    if not isinstance(filas, (list, tuple)):
        return out
    for f in filas:
        if isinstance(f, (list, tuple)):
            out.append([str(c) if c is not None else "" for c in f])
        elif isinstance(f, dict):
            out.append([str(v) if v is not None else "" for v in f.values()])
        elif f is None:
            continue
        else:
            out.append([str(f)])
    return out


def _parse_markdown_table(text: str):
    """Parsea una tabla Markdown → (encabezados, filas). (None, None) si no aplica.

    Los modelos producen tablas Markdown de forma MUCHO más fiable que arrays JSON
    anidados, así que las aceptamos como fuente alterna (campo `markdown`, o una
    tabla Markdown que el modelo metió por error dentro de un párrafo)."""
    import re as _re
    lines = [l.rstrip() for l in (text or "").splitlines() if l.strip()]
    pipe_lines = [l for l in lines if "|" in l]
    if len(pipe_lines) < 2:
        return None, None

    def _is_sep(l: str) -> bool:
        return bool(_re.fullmatch(r"[\s|:\-]+", l)) and "-" in l

    sep_idx = next((i for i, l in enumerate(pipe_lines) if _is_sep(l)), None)
    if not sep_idx:  # None o 0 → no hay header antes del separador
        return None, None

    def _split(l: str) -> list:
        return [c.strip() for c in l.strip().strip("|").split("|")]

    encabezados = _split(pipe_lines[sep_idx - 1])
    filas = [_split(l) for l in pipe_lines[sep_idx + 1:] if not _is_sep(l)]
    return encabezados, filas


def _render_table(doc: Document, encabezados: list, filas: list) -> None:
    cols = max(len(encabezados), max((len(f) for f in filas), default=0))
    if cols == 0:
        return
    rows = (1 if encabezados else 0) + len(filas)
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    r = 0
    if encabezados:
        for c in range(cols):
            cell = table.cell(0, c)
            _cell_shade(cell, PRIMARY)
            cell.text = ""
            run = cell.paragraphs[0].add_run(
                str(encabezados[c]) if c < len(encabezados) else ""
            )
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
        r = 1

    for i, fila in enumerate(filas):
        for c in range(cols):
            cell = table.cell(r + i, c)
            if i % 2 == 1:
                _cell_shade(cell, LIGHT)
            cell.text = ""
            run = cell.paragraphs[0].add_run(
                str(fila[c]) if c < len(fila) else ""
            )
            run.font.size = Pt(10)
    doc.add_paragraph()


def _render_section(doc: Document, block: dict) -> None:
    tipo = (block.get("tipo") or "parrafo").lower()

    if tipo == "heading":
        nivel = block.get("nivel", 1)
        nivel = nivel if nivel in (1, 2, 3) else 1
        p = doc.add_paragraph(block.get("texto", ""), style=f"Heading {nivel}")
        if nivel == 1:
            _bottom_border(p, ACCENT, size=8)
    elif tipo == "parrafo":
        texto = block.get("texto", "")
        # Defensa: si el modelo metió una tabla Markdown dentro de un párrafo,
        # la renderizamos como tabla real en vez de texto plano con pipes.
        enc_md, fil_md = _parse_markdown_table(texto)
        if fil_md and enc_md and len(enc_md) >= 2:
            _render_table(doc, enc_md, _normalize_filas(fil_md))
        else:
            doc.add_paragraph(texto)
    elif tipo == "lista":
        estilo = "List Number" if block.get("estilo") == "numerada" else "List Bullet"
        for item in block.get("items", []):
            try:
                doc.add_paragraph(str(item), style=estilo)
            except KeyError:
                doc.add_paragraph(f"• {item}")
    elif tipo == "tabla":
        encabezados = block.get("encabezados", []) or []
        filas = _normalize_filas(block.get("filas", []))
        # Fallback Markdown: campo `markdown`, o (si no hay filas) una tabla
        # Markdown que el modelo metió en `texto`. Recupera tablas que de otro
        # modo saldrían solo con encabezados (el bug de "filas vacías").
        md = block.get("markdown") or ""
        if not filas and not md:
            md = block.get("texto") or ""
        if md:
            enc_md, fil_md = _parse_markdown_table(md)
            if fil_md:
                if not encabezados and enc_md:
                    encabezados = enc_md
                filas = _normalize_filas(fil_md)
        _render_table(doc, encabezados, filas)
    elif tipo == "cita":
        try:
            doc.add_paragraph(block.get("texto", ""), style="Intense Quote")
        except KeyError:
            p = doc.add_paragraph()
            run = p.add_run(block.get("texto", ""))
            run.font.italic = True
    elif tipo in ("salto", "pagebreak"):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        doc.add_paragraph(block.get("texto", ""))


def _finalize(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── API pública ────────────────────────────────────────────────────────────

def build_docx(spec: dict, *, agente: str = "PipeOs", fecha: datetime | None = None) -> bytes:
    """Construye un .docx universal a partir de una estructura de secciones."""
    fecha = fecha or datetime.now()
    fecha_str = _fecha_larga(fecha)
    titulo = spec.get("titulo") or "Documento"
    subtitulo = spec.get("subtitulo") or ""
    tipo = spec.get("tipo") or "documento"

    doc = Document()
    _apply_theme(doc)
    _build_header(doc)
    _build_footer(doc, agente, fecha_str)
    _build_cover(doc, titulo, subtitulo, agente, tipo, fecha_str)

    for block in spec.get("secciones", []):
        if isinstance(block, dict):
            _render_section(doc, block)

    return _finalize(doc)


def build_tutela(payload: dict, *, agente: str = "PipeOs",
                 fecha: datetime | None = None) -> bytes:
    """Construye una Acción de Tutela (Decreto 2591/1991, art. 14) como BORRADOR.

    Marca de agua diagonal + disclaimer en pie en todas las páginas. NUNCA es un
    documento definitivo: requiere validación y firma de abogado titulado.
    """
    fecha = fecha or datetime.now()
    fecha_str = _fecha_larga(fecha)
    ciudad = payload.get("ciudad") or "_____________"

    doc = Document()
    _apply_theme(doc)
    _build_header(doc)

    section = doc.sections[0]
    _add_watermark(section, "BORRADOR")

    # Pie con disclaimer legal en todas las páginas
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = fp.add_run(
        "BORRADOR — REQUIERE VALIDACIÓN Y FIRMA DE ABOGADO TITULADO · "
        "Generado por PipeOs — no constituye asesoría jurídica definitiva"
    )
    dr.font.size = Pt(7.5)
    dr.font.bold = True
    dr.font.color.rgb = RGBColor.from_string("B00020")

    # Encabezado del juez
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = h.add_run("Señor(a) Juez (Reparto)" if not payload.get("juez_competente")
                   else payload["juez_competente"])
    hr.font.bold = True
    doc.add_paragraph(f"{ciudad}").alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ref = doc.add_paragraph()
    ref.add_run("Referencia: ").bold = True
    ref.add_run("ACCIÓN DE TUTELA (art. 86 C.P. — Decreto 2591/1991)")

    acc = payload.get("accionante", {})
    saludo = doc.add_paragraph()
    saludo.add_run(
        f"{acc.get('nombre', '_____________')}, identificado(a) con "
        f"{acc.get('tipo_id', 'C.C.')} No. {acc.get('num_id', '__________')}, "
        f"domiciliado(a) en {acc.get('domicilio', '_____________')}, "
        "respetuosamente acudo ante su Despacho para interponer ACCIÓN DE TUTELA, "
        "con fundamento en los siguientes:"
    )

    def _heading(txt):
        doc.add_paragraph(txt, style="Heading 2")

    def _numbered(items):
        for it in items or []:
            try:
                doc.add_paragraph(str(it), style="List Number")
            except KeyError:
                doc.add_paragraph(str(it))

    _heading("I. HECHOS")
    _numbered(payload.get("hechos"))

    _heading("II. DERECHOS FUNDAMENTALES VULNERADOS")
    for d in payload.get("derechos_vulnerados", []):
        p = doc.add_paragraph(style="List Bullet")
        art = d.get("articulo_cp")
        if art:
            r = p.add_run(f"Art. {art} C.P. — ")
            r.bold = True
        p.add_run(d.get("descripcion", ""))

    acc_do = payload.get("accionado", {})
    _heading("III. ENTIDAD / PERSONA ACCIONADA")
    pa = doc.add_paragraph()
    pa.add_run(
        f"{acc_do.get('entidad', '_____________')} "
        f"({acc_do.get('naturaleza', 'naturaleza no especificada')}), "
        f"representada por {acc_do.get('representante', '_____________')}, "
        f"con dirección de notificación en "
        f"{acc_do.get('direccion_notificacion', '_____________')}."
    )

    _heading("IV. PRETENSIONES")
    _numbered(payload.get("pretensiones"))

    if payload.get("medida_cautelar"):
        mc = payload["medida_cautelar"]
        _heading("V. MEDIDA PROVISIONAL (art. 7 Decreto 2591/1991)")
        p = doc.add_paragraph()
        p.add_run(mc.get("descripcion", ""))
        if mc.get("justificacion"):
            j = doc.add_paragraph()
            j.add_run("Justificación: ").bold = True
            j.add_run(mc["justificacion"])

    if payload.get("fundamentos_derecho"):
        _heading("VI. FUNDAMENTOS DE DERECHO")
        doc.add_paragraph(payload["fundamentos_derecho"])

    _heading("VII. JURAMENTO")
    jur = doc.add_paragraph()
    if payload.get("juramento_no_otra_tutela"):
        jur.add_run(
            "Bajo la gravedad del juramento manifiesto que no he presentado otra "
            "acción de tutela por los mismos hechos y derechos (art. 38 Decreto 2591/1991)."
        )
    else:
        jur.add_run(
            "[PENDIENTE: declaración bajo juramento de no haber presentado otra "
            "tutela por los mismos hechos — art. 38 Decreto 2591/1991]"
        ).italic = True

    _heading("VIII. ANEXOS")
    if payload.get("anexos"):
        _numbered(payload["anexos"])
    else:
        doc.add_paragraph("[Relacionar pruebas y documentos de soporte]")

    _heading("IX. NOTIFICACIONES")
    notif = doc.add_paragraph()
    notif.add_run(
        f"Accionante: {acc.get('domicilio', '_____________')}"
        + (f" · {acc.get('email')}" if acc.get("email") else "")
        + (f" · {acc.get('telefono')}" if acc.get("telefono") else "")
    )
    doc.add_paragraph(
        f"Accionado: {acc_do.get('direccion_notificacion', '_____________')}"
    )

    doc.add_paragraph()
    cierre = doc.add_paragraph()
    cierre.add_run(f"Atentamente,").bold = False
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.add_run("_______________________________\n")
    firma.add_run(f"{acc.get('nombre', '_____________')}\n")
    firma.add_run(f"{acc.get('tipo_id', 'C.C.')} No. {acc.get('num_id', '__________')}")
    doc.add_paragraph(f"{ciudad}, {fecha_str}")

    return _finalize(doc)
